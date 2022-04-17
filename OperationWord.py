import os


from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn


def addPicAndTitle(doc, title, pics = []):
    doc.add_heading(title, 5)
    print("开始放图片到word")
    
    pics.sort()
    for pic in pics:
        print("正在放{}".format(pic))
        doc.add_picture(pic, width=Inches(5.5))
    print("完成放图片到word")
    doc.save('test.docx')


def opWord(name):
    file = os.path.join(os.path.curdir, name)
    print(file)
    # 判断当前目录是否为文件夹
    if os.path.exists(file) and os.path.isfile(file) :
        # 进行转换
        word = Document(name)
        return word
    else:
        word = Document()
        word.save(name)
        word = Document(name)
        return word



if __name__ == "__main__":

    doc = Document()
    doc.add_heading("title", 0)
    doc.save('test.docx')
