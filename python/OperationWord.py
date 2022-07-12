import os
import sys
import io

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn

# sys.stdout = io.TextIOWrapper(sys.stdout.buffer,
#                               encoding='utf-8')  #改变标准输出的默认编码


def addPicAndTitle(wordPath, wordName, doc, title, pics=[]):
    file = os.path.join(wordPath, wordName)

    doc.add_heading(title, 5)
    print(u"开始放图片到word")

    pics.sort()
    for pic in pics:
        print(u"正在放{}".format(pic))
        doc.add_picture(pic, width=Inches(5.5))
    print(u"完成放图片到word")
    doc.save(file)


def opWord(wordPath, wordName):
    file = os.path.join(wordPath, wordName)
    print(file)
    # 判断当前目录是否为文件夹
    if os.path.exists(file) and os.path.isfile(file):
        # 进行转换
        word = Document(file)
        return word
    else:
        word = Document()
        word.save(file)
        word = Document(file)
        return word


if __name__ == "__main__":

    doc = Document()
    doc.add_heading("title", 0)
    doc.save('test.docx')
